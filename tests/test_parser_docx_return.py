"""
parser_docx 返回值契约测试（TC-D）。

验证 parse_document 始终返回三元组 (md, pages, tables)，
避免 main.py 解包二元组导致 Too many values to unpack（历史 bug）。

注意：DOCX 的表格内容写入 Markdown 的 ``## Tables`` 章节，``tables`` 返回值
始终为空列表（与 PDF 不同，PDF 在 parse_pdf 中返回 tables 列表）。
"""

from pathlib import Path

from docx import Document

from src.parser_docx import parse_document


def _make_docx(path: Path) -> None:
    d = Document()
    d.add_heading("Section 1", level=1)
    d.add_paragraph("U-plane 数据通过 eCPRI 传输。")
    # 一个表格，验证表格内容进入 Markdown
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "参数"
    t.cell(0, 1).text = "值"
    t.cell(1, 0).text = "FFT Size"
    t.cell(1, 1).text = "4096"
    d.save(str(path))


def test_parse_document_returns_three_tuple(tmp_path):
    p = tmp_path / "sample.docx"
    _make_docx(p)
    result = parse_document(str(p))
    assert isinstance(result, tuple), "parse_document 必须返回元组"
    assert len(result) == 3, f"应为三元组 (md, pages, tables)，实际 {len(result)}"
    md, pages, tables = result
    assert isinstance(md, str) and len(md) > 0
    assert isinstance(pages, list)
    assert isinstance(tables, list)
    # DOCX 表格内容写入 Markdown 的 ## Tables 章节
    assert "## Tables" in md, "含表格的 docx 应在 Markdown 中生成 ## Tables 章节"
    assert "FFT Size" in md, "表格单元格内容应出现在 Markdown 中"
    # docx 的 tables 返回空列表（符合设计：表格已并入 Markdown）
    assert tables == [], "DOCX 的 parse_document 返回空 tables 列表（表格已并入 Markdown）"


def test_parse_document_without_tables(tmp_path):
    p = tmp_path / "notable.docx"
    d = Document()
    d.add_heading("Only Text", level=1)
    d.add_paragraph("hello world")
    d.save(str(p))
    md, pages, tables = parse_document(str(p))
    assert isinstance(tables, list)
    assert tables == []
