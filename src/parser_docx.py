"""
Word (docx) 解析模块

将 .docx 文档解析为结构化 Markdown（保留章节层级、段落定位信息）。
"""

import logging
import re
from pathlib import Path
from typing import Optional

logger = logging.getLogger("fh_protocol_compare.parser_docx")


# ---------------------------------------------------------------------------
# 章节标题识别（Word 样式索引映射）
# ---------------------------------------------------------------------------

# Word 内置样式 → Markdown 标题级别
STYLE_TO_LEVEL = {
    "heading 1": 1,
    "heading1": 1,
    "heading 2": 2,
    "heading2": 2,
    "heading 3": 3,
    "heading3": 3,
    "heading 4": 4,
    "heading4": 4,
    "title": 1,
    "caption": 0,  # 跳过
}

# 编号章节行（如 "1.2.3  Signal Flow"）
NUMBERED_TITLE_PATTERNS = [
    re.compile(r"^(\d+(?:\.\d+)*)\s+(.+)$"),
    re.compile(r"^([A-Z])\.\s+(.+)$"),
    re.compile(r"^(Annex|Appendix)\s+([A-Z])\.?\s+(.+)$", re.I),
]

HEADER_FOOTER_PATTERNS = [
    re.compile(r"^3GPP\s+TS\s+\d+\.\d+", re.I),
    re.compile(r"^Page\s+\d+\s+of\s+\d+", re.I),
    re.compile(r"^\d+\s+Release\s+\d+", re.I),
]


def _is_header_footer(text: str) -> bool:
    text = text.strip()
    if not text or len(text) < 3:
        return True
    for pat in HEADER_FOOTER_PATTERNS:
        if pat.match(text):
            return True
    return False


def _parse_numbered_title(text: str) -> Optional[tuple[int, str]]:
    """解析带编号的标题文本"""
    text = text.strip()
    for pat in NUMBERED_TITLE_PATTERNS:
        m = pat.match(text)
        if m:
            groups = m.groups()
            if len(groups) == 2:          # "1.2  Title" 或 "A.  Title"
                if groups[0].upper() in ("A", "B", "C"):
                    return 1, groups[1].strip()
                return 1, groups[1].strip()
            elif len(groups) == 3:        # "Annex A. Scope"
                return 1, groups[2].strip()
            elif len(groups) == 4:
                return 1, groups[3].strip()
    return None


def _style_to_level(style_name: str) -> Optional[int]:
    """将 Word 样式名转换为 Markdown 级别"""
    if not style_name:
        return None
    s = style_name.lower().strip()
    # 精确匹配
    if s in STYLE_TO_LEVEL:
        return STYLE_TO_LEVEL[s]
    # 前缀匹配
    for key, lvl in STYLE_TO_LEVEL.items():
        if s.startswith(key):
            return lvl
    return None


# ---------------------------------------------------------------------------
# 核心解析
# ---------------------------------------------------------------------------

def parse_docx(docx_path: str) -> tuple[str, list[dict]]:
    """
    解析 .docx 文件为结构化 Markdown。

    Args:
        docx_path: .docx 文件路径

    Returns:
        (structured_markdown: str, raw_paragraphs: list[dict])

    raw_paragraphs 元素：
    {
        "para_num": int,
        "style": str,        # Word 样式名
        "level": int,        # Markdown 标题级别（0=正文）
        "text": str,
        "page_num": int,     # Word 页码（可能有偏差）
    }
    """
    import os
    from docx import Document

    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"docx 文件不存在: {docx_path}")

    logger.info(f"[DOCX] 开始解析: {docx_path}")

    doc = Document(docx_path)
    raw_paragraphs = []
    current_h1 = ""
    current_h2 = ""
    current_h3 = ""

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style_name = para.style.name if para.style else ""

        if not text:
            continue

        if _is_header_footer(text):
            continue

        # 判断是否为标题
        level = _style_to_level(style_name)
        is_title = False
        if level is not None and level > 0:
            is_title = True
        else:
            # 尝试从文本内容推断
            parsed = _parse_numbered_title(text)
            if parsed:
                level_hint, _ = parsed
                level = level_hint
                is_title = True

        if is_title:
            level = level or 1
            if level == 1:
                current_h1 = text
                current_h2 = ""
                current_h3 = ""
            elif level == 2:
                current_h2 = text
                current_h3 = ""
            elif level == 3:
                current_h3 = text

            md_prefix = "#" * min(level, 6)
            raw_paragraphs.append({
                "para_num": i,
                "style": style_name,
                "level": level,
                "text": text,
                "page_num": 0,   # python-docx 不直接暴露页码
                "is_title": True,
            })
        else:
            raw_paragraphs.append({
                "para_num": i,
                "style": style_name,
                "level": 0,
                "text": text,
                "page_num": 0,
                "is_title": False,
            })

    # 提取表格
    tables = []
    for i, table in enumerate(doc.tables):
        tbl_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            tbl_data.append(row_data)
        tables.append({"table_index": i, "data": tbl_data})

    # 生成 Markdown
    md_lines = []
    for p in raw_paragraphs:
        if p["is_title"]:
            md_lines.append(f"{'#' * min(p['level'], 6)} {p['text']}\n")
        else:
            md_lines.append(f"{p['text']}\n\n")

    # 追加表格
    for tbl in tables:
        md_lines.append("\n## Tables\n\n")
        for row in tbl["data"]:
            md_lines.append("| " + " | ".join(row) + " |\n")
        md_lines.append("\n")

    md = "".join(md_lines)
    logger.info(f"[DOCX] 解析完成，{len(raw_paragraphs)} 段落，{len(tables)} 表格")
    return md, raw_paragraphs


def parse_document(doc_path: str) -> tuple[str, list[dict]]:
    """
    统一入口：根据文件扩展名自动选择解析器。

    Args:
        doc_path: PDF 或 DOCX 文件路径

    Returns:
        (structured_markdown: str, raw_pages/paragraphs: list[dict])
    """
    ext = Path(doc_path).suffix.lower()
    if ext == ".pdf":
        from src.parser_pdf import parse_pdf
        return parse_pdf(doc_path)
    elif ext in (".docx", ".doc"):
        return parse_docx(doc_path)
    else:
        raise ValueError(f"不支持的文档格式: {ext}，仅支持 PDF 和 DOCX")


# ---------------------------------------------------------------------------
# 调试入口
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys
    import logging

    logging.basicConfig(level=logging.DEBUG, format="%(levelname)s | %(message)s")
    if len(sys.argv) < 2:
        print("用法: python -m src.parser_docx <docx_path>")
        sys.exit(1)

    md, paras = parse_docx(sys.argv[1])
    print(md[:3000])
